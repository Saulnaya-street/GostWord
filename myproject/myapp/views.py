from django.http import HttpResponse
from django.views.decorators.http import require_POST
from docx import Document
from io import BytesIO
import matplotlib.pyplot as plt
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from django.contrib.auth import login
from django.contrib.auth.forms import AuthenticationForm
from django.shortcuts import render, redirect
from django.urls import reverse
from django.contrib.auth.decorators import login_required
from django.http import HttpResponseRedirect
from django.contrib.auth import logout



@login_required
def confirm_logout(request):
    if request.method == 'POST':
        logout(request)
        return HttpResponseRedirect(reverse('home'))
    return render(request, 'confirm_logout.html')

def register_view(request):
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)  # Автоматический вход пользователя в систему
            return redirect('profile')  # Перенаправление на страницу профиля после регистрации
    else:
        form = CustomUserCreationForm()
    return render(request, 'register.html', {'form': form})

def login_view(request):
    if request.method == 'POST':
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            return redirect(reverse('profile'))  # Предполагается, что у вас есть URL с именем 'home'
    else:
        form = AuthenticationForm()
    return render(request, 'login.html', {'form': form})

@login_required
def profile_view(request):
    return render(request, 'profile.html', {'user': request.user})

def home(request):
    # Эта функция отображает главную страницу.
    return render(request, 'home.html')

def about(request):
    # Логика для страницы "О приложении"
    return render(request, 'about.html')

def contact(request):
    # Логика для страницы "Контакты"
    return render(request, 'contact.html')

def terms_view(request):
    return render(request, 'terms.html')

def privacy_policy_view(request):
    return render(request, 'privacy_policy.html')

def add_elements(request):
    return render(request, 'add.html')




def set_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            # Установка границ для всех ячеек
            borders = OxmlElement('w:tcBorders')
            for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')  # размер в восьмых точке (1/8 pt)
                border.set(qn('w:space'), '0')  # без пробелов между границами
                border.set(qn('w:color'), 'auto')
                borders.append(border)

            tcPr.append(borders)


def latex_to_image(latex_str):
    # Настройка matplotlib для использования без X-server (для серверных сценариев)
    plt.switch_backend('Agg')

    # Создаем изображение с LaTeX формулой
    fig = plt.figure()
    text = fig.text(0, 0, f'${latex_str}$', fontsize=14)

    # Автоматически определяем размер бокса вокруг текста
    fig.savefig(BytesIO(), format='png')  # Нам нужен только размер бокса, изображение не сохраняется
    bbox = text.get_window_extent()

    # Задаем размер изображения, равный размеру бокса, и рисуем текст
    fig.set_size_inches(bbox.width / fig.dpi, bbox.height / fig.dpi)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.axis('off')
    ax.text(0, 0, f'${latex_str}$', fontsize=14, ha='left', va='bottom')

    # Сохранение в BytesIO объект
    img_stream = BytesIO()
    plt.savefig(img_stream, format='png')
    img_stream.seek(0)
    plt.close(fig)

    return img_stream

@require_POST  # Этот декоратор гарантирует, что эта функция вызывается только для POST-запросов.
def create_word_document(request):
    document = Document()
    element_count = int(request.POST.get('elementCount', 0))

    for i in range(element_count):
        element_type = request.POST.get(f'elementType{i}')
        if element_type == 'text':
            text = request.POST.get(f'text{i}', '')
            document.add_paragraph(text)
        elif element_type == 'header':
            header = request.POST.get(f'header{i}', '')
            document.add_heading(header, level=2)
        elif element_type == 'formula':
            # Получаем LaTeX формулу из POST-запроса
            formula = request.POST.get(f'formula{i}', '')
            # Преобразуем LaTeX формулу в изображение
            image_stream = latex_to_image(formula)  # Получаем объект BytesIO с изображением
            image_stream.seek(0)  # Сбрасываем указатель файла на начало, это важно
            # Добавляем изображение в документ
            paragraph = document.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(image_stream, width=Inches(5))  # Добавляем изображение, используя объект BytesIO
        elif element_type == 'image':
            image_file = request.FILES.get(f'image{i}')  # Получаем файл изображения
            if image_file:
                # Сохраняем изображение во временный BytesIO объект
                image_stream = BytesIO(image_file.read())
                image_stream.seek(0)  # Сброс позиции потока в начало
                paragraph = document.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(image_stream, width=Inches(5))  # Добавление  изображение в документ
        elif element_type == 'table':
            rows = int(request.POST.get(f'tableRows{i}'))
            cols = int(request.POST.get(f'tableCols{i}'))
            table = document.add_table(rows=rows, cols=cols, style='Table Grid')  #  стиль 'Table Grid'

                # Применяем границы к таблице
            set_table_borders(table)

            for r in range(rows):
                for c in range(cols):
                    cell_value = request.POST.get(f'cell{i}_{r}_{c}', '')
                    table.cell(r, c).text = cell_value

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    filename = "Generated_Document.docx"
    response = HttpResponse(
        file_stream.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response
